import zipfile, os, sys, re, requests
from pathlib import Path
from bs4 import BeautifulSoup
# from google.cloud import translate
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Length, Pt
from io import StringIO
from flask import send_file
import json
import io 

# os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r"/Users/dannijiang/Documents/programming projects/playpython/secret-walker.json"
# translator = translate.TranslationServiceClient()
# parent = "projects/secret-walker-328106/locations/global"

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def extract_to_doc(f,url):
    
    def index_and_translate(url):
        # Scraping the playlist for video titles
        playlist_html=requests.get(url)
        playlist_soup=BeautifulSoup(playlist_html.text,'html.parser')
        playlist_titles_tagged=playlist_soup.find_all(class_=['title', 'playlist-label'])

        # Open template document to populate
        ordered_links={}

        for playlist in playlist_titles_tagged:
            if playlist == playlist_titles_tagged[0]: # Remove the first heading
                continue
            p_text=playlist.text.strip()
            # if turn_on_mandarin == 'y':   # Mandarin translation option
            #     title_in_arrays=[p_text]
            #     title_translated=translator.translate_text(title_in_arrays,target_language_code='zh-CN', parent=parent, source_language_code='en').translations[0].translated_text
            # else:
            title_translated=p_text

            if p_text in extracted_links: # Hyperlinking titles with links
                hyperlink = add_hyperlink(p, extracted_links[p_text], title_translated, '3755DF', False)
                p.add_run().add_break()
            else:   # Adding chapter headings
                p.add_run().add_break()
                heading=p.add_run(title_translated)
                heading.font.bold=True
                p.add_run().add_break()


    frc_lab_animal ={'Basics of Animal Handling': '动物处理的基础知识', 'Rodent Handling and Restraint Techniques': '啮齿动物处理和限制技术', 'An Introduction to the Laboratory Mouse: Mus musculus': '实验室小鼠简介：Mus musculus', 'Basic Mouse Care and Maintenance': '基本鼠标保养和维护', 'Basic Care Procedures': '基本护理程序', 'Rodent Identification I': '啮齿动物鉴定 I', 'Tickling, a Technique for Inducing Positive Affect When Handling Rats': '挠痒痒，一种在处理大鼠时诱导积极影响的技术', 'A Protocol for Using Gene Set Enrichment Analysis to Identify the Appropriate Animal Model for Translational Research': '使用基因集富集分析来确定适用于转化研究的动物模型的协议', 'Animal Feeding and Administration': '动物饲养和管理', 'Compound Administration I': '药物注射 I', 'Compound Administration II': '药物注射 II', 'Compound Administration III': '药物注射 III', 'Compound Administration IV': '化合物给药 IV', 'Manual Restraint and Common Compound Administration Routes in Mice and Rats': '小鼠和大鼠的手动约束和常见化合物给药途径', 'Monitoring Vitals': '监测生命体征', 'Systematic Assessment of Well-Being in Mice for Procedures Using General Anesthesia': '使用全身麻醉对小鼠的健康状况进行系统评估', 'Implantation of Radiotelemetry Transmitters Yielding Data on ECG, Heart Rate, Core Body Temperature and Activity in Free-moving Laboratory Mice': '植入无线电遥测发射器产生关于自由移动实验室小鼠心电图、心率、核心体温和活动的数据', 'Echocardiographic and Histological Examination of Cardiac Morphology in the Mouse': '小鼠心脏形态的超声心动图和组织学检查', 'Breeding': '配种', 'Fundamentals of Breeding and Weaning': '育种和断奶的基本原理', 'Development and Reproduction of the Laboratory Mouse': '实验室小鼠的开发和繁殖', 'Assessment of Sexual Behavior of Male Mice': '雄性小鼠性行为的评估', 'Surgery and Invasive Procedures': '手术和侵入性手术', 'Blood Withdrawal I': '抽血Ⅰ', 'Blood Withdrawal II': '抽血II', 'Anesthesia Induction and Maintenance': '麻醉诱导和维持', 'Considerations for Rodent Surgery': '啮齿动物手术的注意事项', 'Principles of Rodent Surgery for the New Surgeon': '新外科医生的啮齿动物手术原则', 'Tissue Harvest': '组织采集', 'Whole Animal Perfusion Fixation for Rodents': '啮齿类动物全动物灌注固定', 'Diagnostic Necropsy and Tissue Harvest': '诊断性尸检和组织采集', 'Sterile Tissue Harvest': '无菌组织采集', 'Diagnostic Necropsy and Selected Tissue and Sample Collection in Rats and Mice': '大鼠和小鼠的诊断性尸检和选择的组织和样本收集', 'Techniques in Rodent Models': '啮齿动物模型技术', 'Mouse Genotyping': '小鼠基因分型', 'Introducing Experimental Agents into the Mouse': '将实验试剂引入小鼠', 'Rodent Stereotaxic Surgery': '啮齿动物立体定向手术', 'Murine In Utero Electroporation': '小鼠子宫内电穿孔', 'Applying Stereotactic Injection Technique to Study Genetic Effects on Animal Behaviors': '应用立体定向注射技术研究遗传对动物行为的影响', 'A Low Cost Setup for Behavioral Audiometry in Rodents': '啮齿类动物行为测听的低成本设置', 'Techniques in Xenopus Models': '非洲爪蟾模型中的技术', 'Reproductive Techniques for Ovarian Monitoring and Control in Amphibians': '两栖动物卵巢监测和控制的生殖技术', 'Fertilization of Xenopus oocytes using the Host Transfer Method': '使用宿主转移方法对爪蟾卵母细胞进行受精', 'The Xenopus Oocyte Cut-open Vaseline Gap Voltage-clamp Technique With Fluorometry': '爪蟾卵母细胞切开凡士林间隙电压钳技术与荧光法', 'Techniques in Rabbit Models': '兔子模型中的技术', 'Methods of Pairing and Pair Maintenance of New Zealand White Rabbits (Oryctolagus Cuniculus) Via Behavioral Ethogram, Monitoring, and Interventions': '通过行为 Ethogram、监测和干预对新西兰白兔 (Oryctolagus Cuniculus) 进行配对和配对维护的方法', 'Transthoracic Echocardiographic Examination in the Rabbit Model': '兔模型经胸超声心动图检查'}
    cell_bio={'1. 绪论  细胞的结构与功能概述': '1. 绪论 细胞的结构与功能概述', 'What is Evolutionary History?': '什么是进化史？', 'The Tree of Life - Bacteria, Archaea, Eukaryotes': '生命之树 - 细菌、古细菌、真核生物', 'Levels of Organization': '组织层次', 'What are Cells?': '什么是细胞？', 'Cell Size': '单元格大小', 'Eukaryotic Compartmentalization': '真核区室化', 'Prokaryotic Cells': '原核细胞', 'Cytoplasm': '细胞质', 'Tissues': '纸巾', 'Cell Structure- Concept': '细胞结构-概念', 'Cell Structure - Student Protocol': '细胞结构 - 学生协议', '2. 细胞的分子与物质基础': '2. 细胞的分子与物质基础', 'Chemistry of Carbohydrates': '碳水化合物化学', 'Dehydration Synthesis': '脱水合成', 'Hydrolysis': '水解', 'What are Lipids?': '什么是脂质？', 'Structure of Lipids': '脂质的结构', 'Polymers': '聚合物', 'Protein Organization': '蛋白质组织', 'Protein Folding': '蛋白质折叠', 'What are Nucleic Acids?': '什么是核酸？', 'Phosphodiester Linkages': '磷酸二酯键', 'Noncovalent Attractions in Biomolecules': '生物分子中的非共价吸引力', 'Macromolecules- Concept': '大分子-概念', 'Macromolecules - Student Protocol': '大分子 - 学生协议', '3. 细胞生物学的研究方法': '3. 细胞生物学的研究方法', 'Magnetic Activated Cell Sorting (MACS): Isolation of Thymic T Lymphocytes': '磁激活细胞分选 (MACS)：胸腺 T 淋巴细胞的分离', 'Flow Cytometry and Fluorescence-Activated Cell Sorting (FACS): Isolation of Splenic B Lymphocytes': '流式细胞术和荧光激活细胞分选 (FACS)：脾 B 淋巴细胞的分离', 'An Introduction to the Centrifuge': '离心机简介', 'Density Gradient Ultracentrifugation': '密度梯度超速离心', '(1) 显微镜技术': '(1)显微镜技术', 'Introduction to Light Microscopy': '光学显微镜简介', 'Proper Care and Cleaning of the Microscope': '正确保养和清洁显微镜', 'Introduction to Fluorescence Microscopy': '荧光显微镜简介', 'Imaging Biological Samples with Optical and Confocal Microscopy': '使用光学和共焦显微镜对生物样品进行成像', 'Histological Sample Preparation for Light Microscopy': '用于光学显微镜的组织学样品制备', '(2) 细胞的分离和培养': '(2) 细胞的分离和培养', 'Using a Hemacytometer to Count Cells': '使用血细胞计数器计数细胞', 'Passaging Cells': '传代细胞', 'Trypsinizing and Subculturing Mammalian Cells': '哺乳动物细胞的胰蛋白酶消化和传代培养', 'Primary Neuronal Cultures': '原代神经元培养物', 'An Introduction to Transfection': '转染简介', '(3) 细胞组分的分离和纯化技术': '(3) 细胞单体的分离和精炼技术', 'Column Chromatography': '柱色谱', 'Chromatography-based Biomolecule Purification Methods': '基于色谱的生物分子纯化方法', 'Co-Immunoprecipitation and Pull-Down Assays': '免疫共沉淀和下拉测定', 'Immunoprecipitation-Based Techniques: Purification of Endogenous Proteins Using Agarose Beads': '基于免疫沉淀的技术：使用琼脂糖珠纯化内源性蛋白质', 'Separating Protein with SDS-PAGE': '使用 SDS-PAGE 分离蛋白质', 'Gel Purification': '凝胶纯化', '(4) 细胞化学和细胞内分子示踪技术': '(4) 细胞化学和细胞内分子示踪技术', 'Autoradiography as a Simple and Powerful Method for Visualization and Characterization of Pharmacological Targets': '放射自显影作为一种简单而强大的药理靶标可视化和表征方法', 'Revealing Neural Circuit Topography in Multi-Color': '在多色中显示神经电路拓扑', 'Conditional Genetic Transsynaptic Tracing in the Embryonic Mouse Brain': '胚胎小鼠大脑中的条件遗传跨突触追踪', '(5) 生物大分子的结构测定': '(5) 生物大分子的结构补', 'X-ray Diffraction': 'X射线衍射', 'Nuclear Magnetic Resonance (NMR) Spectroscopy': '核磁共振 (NMR) 光谱', 'fMRI: Functional Magnetic Resonance Imaging': 'fMRI：功能性磁共振成像', '4. 细胞膜与物质的跨膜运输': '4. 细胞膜与物质的跨膜运输', 'Membrane Fluidity': '膜流动性', 'The Fluid Mosaic Model': '流体镶嵌模型', 'Protein Associations': '蛋白质关联', 'What is an Electrochemical Gradient?': '什么是电化学梯度？', 'Diffusion and Osmosis- Concept': '扩散和渗透 - 概念', 'Primary Active Transport': '主要主动运输', 'Secondary Active Transport': '次要主动运输', 'Receptor-mediated Endocytosis': '受体介导的内吞作用', 'Pinocytosis': '胞饮作用', 'Phagocytosis': '吞噬作用', 'Exocytosis': '胞吐作用', '5. 细胞的内膜系统、囊泡转运和蛋白分选': '5.细胞的内膜系统、囊泡转运和蛋白分选', 'The Nucleus': '核心', 'Endoplasmic Reticulum': '内质网', 'Ribosomes': '核糖体', 'Reconstitution of Membrane Proteins': '膜蛋白的重建', 'Golgi Apparatus': '高尔基体', 'Mitochondria': '线粒体', '6. 线粒体与细胞能量转换': '6.线粒体与细胞能量转换', 'Primary Production': '初级生产', 'What is Cellular Respiration?': '什么是细胞呼吸？', 'Cellular Respiration- Concept': '细胞呼吸-概念', 'Cellular Respiration - Student Protocol': '细胞呼吸 - 学生协议', 'Electron Transport Chains': '电子传输链', 'Electron Carriers': '电子载体', 'Outcomes of Glycolysis': '糖酵解的结果', '7. 细胞骨架与细胞的运动': '7. 细胞与细胞的运动', 'Microtubules': '微管', 'Cell Migration': '细胞迁移', 'What is Glycolysis?': '什么是糖酵解？', 'An Introduction to Cell Motility and Migration': '细胞运动和迁移简介', 'The Transwell Migration Assay': 'Transwell 迁移分析', 'In vitro Cell Migration and Invasion Assays': '体外细胞迁移和侵袭试验', '8.细胞核': '8.细胞核', 'The Nucleolus': '核仁', 'Chromosome Structure': '染色体结构', 'Chromatin Packaging': '染色质包装', 'Nondisjunction': '不分离', 'Euchromatin': '常染色质', 'Heterochromatin': '异染色质', 'Telomeres and Telomerase': '端粒和端粒酶', 'Polytene Chromosomes': '多线染色体', 'X and Y Chromosomes': 'X 和 Y 染色体', 'Condensins': '凝聚素', 'Cohesins': '粘连蛋白', 'Spreading of Chromatin Modifications': '染色质修饰的传播', 'Duplication of Chromatin Structure': '染色质结构的重复', 'Tumor Progression': '肿瘤进展', 'Loss of Tumor Suppressor Gene Functions': '肿瘤抑制基因功能丧失', 'Abnormal Proliferation': '异常增殖', '9. 细胞连接与细胞粘连 10. 细胞外基质及其与细胞的相互作用': '9. 细胞连接与细胞连 10. 细胞外表面及其与细胞的粘连', 'Gap Junctions': '缝隙连接', 'The Extracellular Matrix': '细胞外基质', 'Invasion Assay Using 3D Matrices': '使用 3D 矩阵进行入侵检测', 'The Tumor Microenvironment': '肿瘤微环境', 'Metastasis': '转移', 'Adaptive Mechanisms in Cancer Cells': '癌细胞的适应性机制', '11. 细胞的信号转导': '11.细胞的信号转导', 'Contact-dependent Signaling': '依赖于接触的信令', 'Autocrine Signaling': '自分泌信号', 'Paracrine Signaling': '旁分泌信号', 'G-protein Coupled Receptors': 'G蛋白偶联受体', 'Endocrine Signaling': '内分泌信号', 'Ion Channels': '离子通道', 'What are Second Messengers?': '什么是第二使者？', 'Enzyme-linked Receptors': '酶联受体', '12. 细胞分裂与细胞周期': '12. 细胞分裂与细胞周期', 'Binary Fission': '二分裂', 'Interphase': '相间', 'Mitosis and Cytokinesis': '有丝分裂和细胞分裂', 'Positive Regulator Molecules': '正调节分子', 'Negative Regulator Molecules': '负调节分子', 'Live Cell Imaging of Mitosis': '有丝分裂的活细胞成像', 'Meiosis I': '减数分裂Ⅰ', 'Meiosis II': '减数分裂II', 'Crossing Over': '穿越', 'Cancer-Critical Genes I: Proto-oncogenes': '癌症关键基因 I：原癌基因', 'Cancer-Critical Genes II: Tumor Suppressor Genes': '癌症关键基因 II：肿瘤抑制基因', '13.  细胞分化': '13.细胞神经', 'Embryonic Stem Cell Culture and Differentiation': '胚胎干细胞培养和分化', 'Induced Pluripotency': '诱导多能性', 'Fate Mapping': '命运映射', 'Cancer': '癌症', 'Using Mouse Mammary Tumor Cells to Teach Core Biology Concepts: A Simple Lab Module': '使用小鼠乳腺肿瘤细胞教授核心生物学概念：一个简单的实验室模块', '14. 细胞衰老和死亡': '14.细胞衰老和死亡', 'An Introduction to Aging and Regeneration': '老化与再生简介', 'An Introduction to Cell Death': '细胞死亡简介', 'Tissue Regeneration with Somatic Stem Cells': '用体干细胞进行组织再生', '15. 干细胞与组织的维持与再生': '15. 干细胞与组织的生存与再生', 'An Introduction to Stem Cell Biology': '干细胞生物学导论', 'Adult Stem Cells': '成体干细胞', 'Induced Pluripotent Stem Cells': '诱导多能干细胞', 'Cancer Stem Cells and Tumor Maintenance': '癌症干细胞和肿瘤维持'}
    translated_arrays=[cell_bio, frc_lab_animal]

    doc = docx.Document('templates/template.docx')

    zipfile_ob = zipfile.ZipFile(f)
    file_names = zipfile_ob.namelist()
    file_names = [file_name for file_name in file_names if file_name.endswith(".html")]

    extracted_links = {}
    for name in file_names:
        with zipfile_ob.open(name) as fp:
            vid_soup=BeautifulSoup(fp,'html.parser')
            vid_link=vid_soup.iframe['src']
            vid_name=vid_soup.title.text
            pattern=r'[0-9]{1,}\.[0-9]{1,}\s'
            vid_name_clean=re.sub(pattern,'',vid_name)
            extracted_links[vid_name_clean]=vid_link

    print(extracted_links)
    # print('WORKING FILE: '+f_stem)
    for key, value in extracted_links.items():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing=Pt(24)
        playlist_url=url

        
        try: # Mapping to exisitng translations
            playlist_number=int(playlist_url)
            pre_tranlated_playlist=translated_arrays[playlist_number]
            for key, value in pre_tranlated_playlist.items():
                if key in extracted_links: # Hyperlinking titles with links
                    hyperlink = add_hyperlink(p, extracted_links[str(key)], value, '3755DF', False)
                    p.add_run().add_break()
                else:   # Adding chapter headings
                    p.add_run().add_break()
                    heading=p.add_run(value)
                    heading.font.bold=True
                    p.add_run().add_break()
                    print("into try")
        except:
            index_and_translate(playlist_url) 
            # print("into except")
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, attachment_filename='report.docx')