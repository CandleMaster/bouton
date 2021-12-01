import zipfile, os, sys, re, requests
from pathlib import Path
from bs4 import BeautifulSoup
# from google.cloud import translate
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Length, Pt
from io import StringIO
from flask import send_file
import json
import io 

# os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r"/Users/dannijiang/Documents/programming projects/playpython/secret-walker.json"
# translator = translate.TranslationServiceClient()
# parent = "projects/secret-walker-328106/locations/global"

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def extract_to_doc(f,url,translation_option):
    

    def index_and_translate(url):
        # Scraping the playlist for video titles
        playlist_html=requests.get(url)
        playlist_soup=BeautifulSoup(playlist_html.text,'html.parser')
        playlist_titles_tagged=playlist_soup.find_all(class_=['title', 'playlist-label'])

        # Open template document to populate
        ordered_links={}


        turn_on_mandarin=translation_option
        for playlist in playlist_titles_tagged:
            if playlist == playlist_titles_tagged[0]: # Remove the first heading
                continue
            p_text=playlist.text.strip()
            # if turn_on_mandarin == 'y':   # Mandarin translation option
            #     title_in_arrays=[p_text]
            #     title_translated=translator.translate_text(title_in_arrays,target_language_code='zh-CN', parent=parent, source_language_code='en').translations[0].translated_text
            # else:
            title_translated=p_text

            if p_text in extracted_links: # Hyperlinking titles with links
                hyperlink = add_hyperlink(p, extracted_links[p_text], title_translated, '3755DF', False)
                p.add_run().add_break()
            else:   # Adding chapter headings
                p.add_run().add_break()
                heading=p.add_run(title_translated)
                heading.font.bold=True
                p.add_run().add_break()

    CellBio=""

    translated_arrays=[CellBio]
    doc = docx.Document('templates/template.docx')

    zipfile_ob = zipfile.ZipFile(f)
    file_names = zipfile_ob.namelist()
    file_names = [file_name for file_name in file_names if file_name.endswith(".html")]

    extracted_links = {}
    for name in file_names:
        with zipfile_ob.open(name) as fp:
            vid_soup=BeautifulSoup(fp,'html.parser')
            vid_link=vid_soup.iframe['src']
            vid_name=vid_soup.title.text
            pattern=r'[0-9]{1,}\.[0-9]{1,}\s'
            vid_name_clean=re.sub(pattern,'',vid_name)
            extracted_links[vid_name_clean]=vid_link

    print(extracted_links)
    # print('WORKING FILE: '+f_stem)
    for key, value in extracted_links.items():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing=Pt(24)
        playlist_url=url
        try: # Mapping to exisitng translations
            playlist_number=int(playlist_url)
            pre_tranlated_playlist=translated_arrays[playlist_number]
            for key, value in pre_tranlated_playlist.items():
                if key in extracted_links: # Hyperlinking titles with links
                    hyperlink = add_hyperlink(p, extracted_links[str(key)], value, '3755DF', False)
                    p.add_run().add_break()
                else:   # Adding chapter headings
                    p.add_run().add_break()
                    heading=p.add_run(value)
                    heading.font.bold=True
                    p.add_run().add_break()
                    print("into try")
        except:
            index_and_translate(playlist_url) 
            # print("into except")
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, attachment_filename='report.docx')